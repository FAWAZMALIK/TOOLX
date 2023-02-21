import requests
import urllib.parse
from flask import redirect, render_template, request, send_file 
from functools import wraps
import pythoncom
from docx2pdf import convert
from pdf2docx import Converter
import PyPDF2
from fpdf import FPDF
from random import randint
import os
import shutil
from docx import Document
from PIL import Image
import img2pdf
import matplotlib.pyplot as plt
import cv2
import easyocr
from moviepy.video.io.VideoFileClip import VideoFileClip
from moviepy.audio.io.AudioFileClip import AudioFileClip
from pydub import AudioSegment 
from pydub.utils import which
import zipfile
from pylab import rcParams
# from IPython.display import Image
rcParams['figure.figsize'] = 8,16
# import fitz

def is_integer(n):
    try:
        float(n)
    except ValueError:
        return False
    else:
        return float(n).is_integer()


def make_zipfile(output_filename, source_dir):
    os.chdir('..')
    relroot = os.path.abspath(os.path.join(source_dir, os.pardir))
    with zipfile.ZipFile(output_filename, "w", zipfile.ZIP_DEFLATED) as zip:
        for root, dirs, files in os.walk(source_dir):
            # add directory (needed for empty dirs)
            zip.write(root, os.path.relpath(root, relroot))
            for file in files:
                filename = os.path.join(root, file)
                if os.path.isfile(filename): # regular files only
                    arcname = os.path.join(os.path.relpath(root, relroot), file)
                    zip.write(filename, arcname)

def apology(message, code=400):
    """Render message as an apology to user."""
    def escape(s):
        """
        Escape special characters.

        https://github.com/jacebrowning/memegen#special-characters
        """
        for old, new in [("-", "--"), (" ", "-"), ("_", "__"), ("?", "~q"),
                         ("%", "~p"), ("#", "~h"), ("/", "~s"), ("\"", "''")]:
            s = s.replace(old, new)
        return s
    return render_template("apology.html", top=code, bottom=escape(message)), code
# pdftodocx

def pdftodocx(file):
    pdf_file = file
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    output = path + 'Converted To Docx.docx'
    cv = Converter(pdf_file)
    cv.convert(output)
    cv.close
    
    
    return output

def docxtopdf(file):
    docx_file = file
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    output = path + 'Converted To PDF.pdf'
    pythoncom.CoInitialize()
    convert(docx_file,output)
    pythoncom.CoUninitialize ()

    return output

def docxtotext(file):
    text_file = Document(file)
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    output = path + 'Converted To Text.txt'
    txt = ""
    for para in text_file.paragraphs:
        txt += para.text + "\n"
    with open(output,'w') as f:
        f.write(txt)
    return output

def pdftotext(file):
    pdf_file = open(file,'rb')
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    output = path + 'Converted To Text.txt'
    pdfreader=PyPDF2.PdfReader(pdf_file)
    x = len(pdfreader.pages)
    ob=pdfreader.pages[x - 1]
    text=ob.extract_text()
    file1 = open(output,'a')
    file1.writelines(text)
    pdf_file.close()
    file1.close()
    return output

def texttopdf(file):
    text_file = open(file,'r')
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    output = path + 'Converted To PDF.pdf'
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size = 17)
    for i in text_file:
        pdf.cell(200, 10, txt = i,ln = 1, align = 'L')
    pdf.output(output)  
    return output

def texttodocx(file):
    text_file = open(file,'r')
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    output = path + 'Converted To DOCX.docx'
    txt = ""
    for i in text_file:
        txt += i
    document = Document()
    document.add_paragraph(txt)
    document.add_page_break()
    document.save(output)
    return output

def tojpg(image):
    im1 = Image.open(image)
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    output = path + 'Converted To JPG.jpg'
    im1.save(output)
    return output

def topdf(image):
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    output = path + 'Converted To PDF.pdf'
    # storing image path
    img_path = image
 
    # storing pdf path
    pdf_path = output
    
    # opening image
    image = Image.open(img_path)
    
    # converting into chunks using img2pdf
    pdf_bytes = img2pdf.convert(image.filename)
    
    # opening or creating pdf file
    file = open(pdf_path, "wb")
    
    # writing pdf files with chunks
    file.write(pdf_bytes)
    
    # closing image file
    image.close()
    
    # closing pdf file
    file.close()
    return output

def topng(image):
    im1 = Image.open(image)
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    output = path + 'Converted To PNG.png'
    im1.save(output)
    return output

def extext(image):
    # here you can use any other language you want
    reader = easyocr.Reader(['en'])
    # using the read text function generating the text from image
    output = reader.readtext(image)
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    output1 = path + 'Extracted Text.txt'
    file1 = open(output1,'a')
    for i in range(len(output)):
        file1.writelines(str(output[i][1]))
        file1.writelines('  ')
    file1.close()
    return output1

    
def exframe(video):
    videofile = video
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    cap= cv2.VideoCapture(videofile)
    i=0
    main = os.getcwd()
    os.chdir(path)
    while(cap.isOpened()):
        ret, frame = cap.read()
        if ret == False:
            break
        cv2.imwrite('Frame'+str(i)+'.jpg',frame)
        i+=1

    cap.release()
    cv2.destroyAllWindows()
    name = str(tag) + '.zip' 
    make_zipfile( name, str(tag))
    outpt = parent_dir + name
    os.chdir(main)
    
    return outpt
    
    
    
    print('done') 
    
def tobw(video , name):
    videofile = video
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    
    # Load the video
    print(video)
    print(name)
    # pydub.AudioSegment.ffmpeg = 'C:\\Users\\NawazHussain\\Desktop\\ffmpeg-master-latest-win64-gpl\\bin\\ffmpeg.exe'
    
    video = AudioSegment.from_file(video, format="mp4")
    
    # Export the audio
    audio = video.export(path + "output.mp3", format="mp3")
    cap = cv2.VideoCapture(videofile)

    # Get the frames per second (fps) of the video
    fps = cap.get(cv2.CAP_PROP_FPS)

    # Define the codec and create a VideoWriter object
    fourcc = cv2.VideoWriter_fourcc(*'mp4v') 
    out = cv2.VideoWriter(path + 'output.mp4v', fourcc, fps, (int(cap.get(3)), int(cap.get(4))))

    while cap.isOpened():
        # Read a frame from the video
        ret, frame = cap.read()

        if ret:
            # Convert the frame to grayscale
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)

            # Write the frame to the output video
            out.write(gray)

            # Display the frame
            # cv2.imshow('frame', gray)

            if cv2.waitKey(1) & 0xFF == ord('q'):
                break
        else:
            break

    # Release the VideoCapture and VideoWriter objects
    cap.release()
    out.release()

    # Close all windows
    cv2.destroyAllWindows()
    # Load the video and audio
    video1 = VideoFileClip(path + 'output.mp4v')
    audio1 = AudioFileClip(path + 'output.mp3')

    # Concatenate the video and audio
    final_video = video1.set_audio(audio1)

    # Write the final video to disk
    output = path + "Coverted To GraeyScale.mp4"
    final_video.write_videofile(output)
    print('done')
    return 



def tosound(video):
    videofile = video
    parent_dir = "./static/uploads/"
    tag = int(randint(1,10000))
    path = parent_dir + str(tag) + '/'
    os.mkdir(path)
    video = AudioSegment.from_file(video, format="mp4")
    output = path + "Coverted to audio.mp3"
    audio = video.export(output, format="mp3")
    return output

def converter(op1,op2,pathf):
    file = pathf
    formats = {1:'mp4',2:'m4a',3:'wav',4:'wma',5:'aac',6:'flac',7:'mp3'}
    if op1 and op2 in formats:
        parent_dir = "./static/uploads/"
        tag = int(randint(1,10000))
        path = parent_dir + str(tag) + '/'
        os.mkdir(path)
        sound = AudioSegment.from_file(file,format=formats[op1])
        output = path + 'Audio.' + formats[op2]
        sound.export(output ,format=formats[op2])
        return output
        
        


     


    #  pythoncom.CoInitialize()
    #     convert("static\\uploads\\" + filename, "static\\uploads\\" + filename +".pdf")
    #     pythoncom.CoInitialize()
    #     try:
    #         return send_file('./static/uploads/Converted.pdf', as_attachment=True)
    #     except Exception as e:
    #         return str(e)
    


